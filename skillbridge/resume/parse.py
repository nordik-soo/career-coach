"""Resume text extraction.

Three input formats supported: PDF (pdfplumber), DOCX (python-docx), and
plain text. The binary file is never persisted — bytes come in, text comes
out, bytes are discarded. See docs/resume-design.md §3 for the storage
policy.

Returns a ParseResult that always has a `text` field. If extraction failed
or produced no usable text, `parse_warning` carries the reason so the
responder can fall back to "could you paste your resume text?" without
crashing.

Format detection prefers magic bytes (PDF starts with %PDF-, DOCX is a ZIP
that starts with PK\\x03\\x04) over filename extension. A user with a
misnamed file (`.txt` for a PDF) still gets the right parser.

Never raises on malformed input. Library exceptions are caught and
surfaced as parse_warning="parse_failed" so the chat flow stays robust.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from typing import Final

log = logging.getLogger(__name__)

# Hard cap on resume file size. 5 MB covers anything a newcomer would
# realistically upload as a resume; bigger uploads are almost always
# non-resume PDFs (full scanned passports, settlement files, etc.). The
# route layer should reject before reading the bytes, but we double-check.
MAX_RESUME_BYTES: Final[int] = 5 * 1024 * 1024

# Minimum text length we'd count as a successful extraction. A scanned PDF
# without OCR typically returns a handful of stray characters; 80 is a low
# bar but rejects truly empty extractions.
MIN_EXTRACTED_CHARS: Final[int] = 80

# Supported content-type tokens returned in ParseResult.content_type.
CONTENT_TYPE_PDF = "pdf"
CONTENT_TYPE_DOCX = "docx"
CONTENT_TYPE_TXT = "txt"
CONTENT_TYPE_UNKNOWN = "unknown"


@dataclass
class ParseResult:
    """Output of parse_resume().

    `text` is always a string (possibly empty). `parse_warning` carries
    the reason when extraction couldn't produce usable text:
      - "empty_input"       → zero bytes
      - "too_large"         → exceeds MAX_RESUME_BYTES
      - "unsupported_format"→ neither PDF, DOCX, nor TXT magic bytes
      - "no_text"           → extracted but result was empty / too short
                              (usually a scanned PDF without OCR)
      - "parse_failed"      → underlying library raised
    """
    text: str
    filename: str
    content_type: str
    byte_count: int
    parse_warning: str | None = None


def parse_resume(file_bytes: bytes, filename: str | None = None) -> ParseResult:
    """Extract plain text from a resume file's bytes.

    Routes by detected content type (magic bytes preferred over filename
    extension). Never raises on malformed input.
    """
    name = (filename or "resume").strip() or "resume"
    n = len(file_bytes)

    if n == 0:
        return ParseResult(
            text="", filename=name, content_type=CONTENT_TYPE_UNKNOWN,
            byte_count=0, parse_warning="empty_input",
        )

    if n > MAX_RESUME_BYTES:
        return ParseResult(
            text="", filename=name, content_type=CONTENT_TYPE_UNKNOWN,
            byte_count=n, parse_warning="too_large",
        )

    content_type = _detect_format(file_bytes, name)

    try:
        if content_type == CONTENT_TYPE_PDF:
            text = _extract_pdf(file_bytes)
        elif content_type == CONTENT_TYPE_DOCX:
            text = _extract_docx(file_bytes)
        elif content_type == CONTENT_TYPE_TXT:
            text = _extract_txt(file_bytes)
        else:
            return ParseResult(
                text="", filename=name, content_type=CONTENT_TYPE_UNKNOWN,
                byte_count=n, parse_warning="unsupported_format",
            )
    except Exception as e:
        log.warning("resume parse failed for %s (%s): %s", name, content_type, e)
        return ParseResult(
            text="", filename=name, content_type=content_type,
            byte_count=n, parse_warning="parse_failed",
        )

    cleaned = _normalize_text(text)

    if len(cleaned) < MIN_EXTRACTED_CHARS:
        # Most common cause: scanned PDF with no embedded text layer.
        # The route layer turns this into a "could you paste the text?"
        # response — don't crash, just degrade.
        return ParseResult(
            text=cleaned, filename=name, content_type=content_type,
            byte_count=n, parse_warning="no_text",
        )

    return ParseResult(
        text=cleaned, filename=name, content_type=content_type,
        byte_count=n, parse_warning=None,
    )


# =========================================================================
# Format detection — magic bytes preferred, filename as fallback
# =========================================================================
def _detect_format(file_bytes: bytes, filename: str) -> str:
    """Detect format. Magic bytes win; filename extension is a fallback."""
    head = file_bytes[:8]

    # PDF: starts with '%PDF-'
    if head.startswith(b"%PDF-"):
        return CONTENT_TYPE_PDF

    # DOCX / any modern Office format: ZIP container starting with 'PK\x03\x04'.
    # We let the docx library reject if it's a different ZIP variant.
    if head.startswith(b"PK\x03\x04"):
        return CONTENT_TYPE_DOCX

    # Filename-extension fallback.
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return CONTENT_TYPE_PDF
    if lower.endswith(".docx"):
        return CONTENT_TYPE_DOCX
    if lower.endswith((".txt", ".text", ".md")):
        return CONTENT_TYPE_TXT

    # Heuristic: mostly-printable bytes → treat as plain text.
    if _is_probably_text(file_bytes):
        return CONTENT_TYPE_TXT

    return CONTENT_TYPE_UNKNOWN


def _is_probably_text(file_bytes: bytes) -> bool:
    """Sample the first 1KB; True if >90% printable UTF-8."""
    sample = file_bytes[:1024]
    try:
        decoded = sample.decode("utf-8")
    except UnicodeDecodeError:
        return False
    if not decoded:
        return False
    printable = sum(1 for c in decoded if c.isprintable() or c in "\n\r\t")
    return printable / len(decoded) > 0.90


# =========================================================================
# Per-format extractors
# =========================================================================
def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber. Returns concatenated pages.

    Lazy import — keeps the package importable even when pdfplumber isn't
    installed (useful for tests that don't exercise PDF parsing).
    """
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx.

    Includes both paragraphs AND table cell text — newcomer resumes often
    use table-based layouts (skill grids, two-column formats) and missing
    table content would lose half the resume.
    """
    import docx  # python-docx ships as the `docx` import name

    document = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_txt(file_bytes: bytes) -> str:
    """Decode plain text. UTF-8 first, latin-1 fallback (always succeeds)."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


# =========================================================================
# Light normalization
# =========================================================================
_WHITESPACE_PATTERN = re.compile(r"[ \t]+")
_BLANK_LINE_RUN = re.compile(r"\n\s*\n\s*\n+")


def _normalize_text(text: str) -> str:
    """Light cleanup. Preserves single newlines so section structure
    survives for the downstream extractor.
    """
    if not text:
        return ""
    # Strip control chars except newline/tab.
    text = "".join(c for c in text if c.isprintable() or c in "\n\r\t")
    # Collapse runs of spaces and tabs (but not newlines).
    text = _WHITESPACE_PATTERN.sub(" ", text)
    # Collapse 3+ consecutive blank lines down to 2.
    text = _BLANK_LINE_RUN.sub("\n\n", text)
    return text.strip()
